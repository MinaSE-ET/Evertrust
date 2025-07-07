from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.pdfgen import canvas
from io import BytesIO
from django.template.loader import render_to_string
from django.utils import timezone
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx2pdf import convert
import tempfile
import json

def generate_certificate_pdf(certificate):
    """Generate PDF certificate using HTML template with WeasyPrint"""
    try:
        # Use WeasyPrint for better HTML to PDF conversion with background support
        from weasyprint import HTML, CSS
        from weasyprint.text.fonts import FontConfiguration
        from django.conf import settings
        from django.templatetags.static import static
        import os
        
        # Render the HTML template
        html_content = render_to_string('certificates/certificate_pdf.html', {
            'certificate': certificate
        })
        
        # Configure WeasyPrint with proper font configuration
        font_config = FontConfiguration()
        
        # Create CSS for better PDF rendering
        css_content = """
        @page {
            size: A4;
            margin: 0;
        }
        
        body {
            font-family: 'Times New Roman', serif;
            line-height: 1.6;
            color: #333;
        }
        
        .pdf-container {
            width: 210mm;
            height: 297mm;
            margin: 0 auto;
            background: rgba(255, 255, 255, 0.95);
        }
        
        .pdf-content {
            padding: 40px;
        }
        
        .pdf-header {
            text-align: center;
            margin-bottom: 30px;
        }
        
        .pdf-title {
            font-size: 2.5rem;
            font-weight: bold;
            color: #1e3a8a;
            margin: 0 0 10px 0;
            text-transform: uppercase;
            letter-spacing: 2px;
        }
        
        .pdf-subtitle {
            font-size: 1.2rem;
            color: #666;
            margin: 0;
        }
        
        .certificate-info {
            background: rgba(255, 255, 255, 0.8);
            padding: 20px;
            border-radius: 8px;
            margin-bottom: 30px;
            border: 2px solid #1e3a8a;
        }
        
        .policy-number {
            font-size: 1.5rem;
            font-weight: bold;
            color: #1e3a8a;
            text-align: center;
            margin-bottom: 15px;
        }
        
        .issue-date {
            text-align: center;
            color: #666;
            font-size: 1rem;
        }
        
        .content-grid {
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 30px;
        }
        
        .info-section {
            background: rgba(255, 255, 255, 0.9);
            padding: 20px;
            border-radius: 8px;
            border: 1px solid #e5e7eb;
        }
        
        .section-title {
            font-size: 1.2rem;
            font-weight: bold;
            color: #1e3a8a;
            margin: 0 0 15px 0;
            text-transform: uppercase;
            letter-spacing: 0.5px;
            border-bottom: 2px solid #1e3a8a;
            padding-bottom: 5px;
        }
        
        .info-row {
            display: flex;
            justify-content: space-between;
            margin-bottom: 8px;
            padding: 4px 0;
            border-bottom: 1px solid #f0f0f0;
        }
        
        .info-row:last-child {
            border-bottom: none;
        }
        
        .label {
            font-weight: bold;
            color: #374151;
            min-width: 120px;
        }
        
        .value {
            font-weight: 500;
            color: #111827;
            text-align: right;
        }
        
        .amount {
            font-family: 'Courier New', monospace;
            font-weight: bold;
            color: #059669;
        }
        
        .signature-seal-section {
            display: flex;
            justify-content: space-between;
            align-items: flex-end;
            margin-top: 30px;
            padding: 20px;
            background: rgba(255, 255, 255, 0.9);
            border-radius: 8px;
        }
        
        .signature-section, .seal-section {
            text-align: center;
            flex: 1;
        }
        
        .signature-img, .seal-img {
            max-width: 150px;
            max-height: 60px;
            display: block;
            margin: 0 auto 10px;
        }
        
        .seal-img {
            max-width: 100px;
            max-height: 100px;
        }
        
        .sig-line {
            width: 150px;
            height: 2px;
            background: #333;
            margin: 10px auto;
        }
        
        .sig-label, .seal-label {
            font-weight: bold;
            color: #374151;
            font-size: 0.9rem;
        }
        
        .qr-section {
            position: absolute;
            top: 20px;
            right: 20px;
            text-align: center;
            background: rgba(255, 255, 255, 0.9);
            padding: 10px;
            border-radius: 8px;
            border: 1px solid #e5e7eb;
        }
        
        .qr-img {
            width: 60px;
            height: 60px;
            display: block;
            margin: 0 auto 5px;
        }
        
        .qr-text {
            font-size: 0.7rem;
            color: #666;
            margin: 0;
            font-weight: bold;
        }
        
        .debit-note {
            page-break-before: always;
            background: rgba(255, 255, 255, 0.95);
            padding: 40px;
            border: 2px solid #1e3a8a;
            border-radius: 8px;
            margin: 20px;
        }
        
        .debit-note-header {
            display: flex;
            align-items: center;
            margin-bottom: 30px;
            padding-bottom: 20px;
            border-bottom: 2px solid #1e3a8a;
        }
        
        .debit-note-logo {
            max-height: 40px;
            margin-right: 20px;
        }
        
        .debit-note-title {
            font-size: 1.5rem;
            font-weight: bold;
            color: #1e3a8a;
            text-transform: uppercase;
            letter-spacing: 1px;
        }
        
        .debit-note-table {
            width: 100%;
            margin-bottom: 30px;
            border-collapse: collapse;
        }
        
        .debit-note-table td {
            padding: 12px;
            border-bottom: 1px solid #e5e7eb;
        }
        
        .debit-note-table .label {
            font-weight: bold;
            color: #374151;
            width: 40%;
        }
        
        .debit-note-table .value {
            color: #111827;
            text-align: left;
        }
        
        .debit-note-amount {
            text-align: center;
            margin: 30px 0;
            padding: 20px;
            background: #f8fafc;
            border-radius: 8px;
            border: 2px solid #1e3a8a;
        }
        
        .amount-label {
            font-weight: bold;
            color: #374151;
            font-size: 1.1rem;
            display: block;
            margin-bottom: 10px;
        }
        
        .amount-value {
            font-size: 2rem;
            font-weight: bold;
            color: #1e3a8a;
            font-family: 'Courier New', monospace;
        }
        
        .debit-note-details {
            margin-top: 20px;
            padding: 20px;
            background: #f9fafb;
            border-radius: 8px;
            border-left: 4px solid #1e3a8a;
        }
        
        .debit-note-details strong {
            color: #1e3a8a;
            display: block;
            margin-bottom: 10px;
        }
        
        .debit-note-details p {
            margin: 0;
            line-height: 1.6;
            color: #374151;
        }
        """
        
        # Create HTML object with proper base URL for static files
        base_url = f"file://{settings.STATIC_ROOT}" if settings.STATIC_ROOT else None
        html = HTML(string=html_content, base_url=base_url)
        
        # Create CSS object
        css = CSS(string=css_content, font_config=font_config)
        
        # Generate PDF
        pdf_content = html.write_pdf(stylesheets=[css], font_config=font_config)
        
        return pdf_content
        
    except ImportError as e:
        print(f"WeasyPrint not available: {e}")
        # Fallback to ReportLab if WeasyPrint is not available
        return generate_certificate_pdf_reportlab(certificate)
    except Exception as e:
        print(f"WeasyPrint PDF generation failed: {e}")
        # Fallback to ReportLab
        return generate_certificate_pdf_reportlab(certificate)

def generate_certificate_pdf_reportlab(certificate):
    """Generate PDF certificate using ReportLab (fallback method)"""
    
    # Create buffer for PDF
    buffer = BytesIO()
    
    # Create PDF document
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    story = []
    
    # Get styles
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=30,
        alignment=TA_CENTER,
        textColor=colors.darkblue
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=12,
        spaceBefore=20,
        textColor=colors.darkblue
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=6
    )
    
    # Header
    story.append(Paragraph("EVER TRUST", title_style))
    story.append(Paragraph(certificate.certificate_title, heading_style))
    story.append(Spacer(1, 20))
    
    # Policy Number
    story.append(Paragraph(f"<b>Policy Number:</b> {certificate.policy_number}", normal_style))
    story.append(Paragraph(f"<b>Issue Date:</b> {certificate.issue_date.strftime('%B %d, %Y')}", normal_style))
    story.append(Spacer(1, 20))
    
    # Client Information
    story.append(Paragraph("CLIENT INFORMATION", heading_style))
    story.append(Paragraph(f"<b>Name:</b> {certificate.client_name}", normal_style))
    story.append(Paragraph(f"<b>Address:</b> {certificate.client_address}", normal_style))
    story.append(Paragraph(f"<b>Phone:</b> {certificate.client_phone}", normal_style))
    if certificate.client_email:
        story.append(Paragraph(f"<b>Email:</b> {certificate.client_email}", normal_style))
    story.append(Paragraph(f"<b>ID Number:</b> {certificate.client_id_number}", normal_style))
    story.append(Paragraph(f"<b>Date of Birth:</b> {certificate.client_date_of_birth.strftime('%B %d, %Y')}", normal_style))
    if certificate.client_bank_name:
        story.append(Paragraph(f"<b>Bank Name:</b> {certificate.client_bank_name}", normal_style))
    if certificate.client_bank_account:
        story.append(Paragraph(f"<b>Bank Account:</b> {certificate.client_bank_account}", normal_style))
    
    # Insurance Details
    story.append(Paragraph("INSURANCE DETAILS", heading_style))
    story.append(Paragraph(f"<b>Insurance Type:</b> {certificate.get_certificate_type_display()}", normal_style))
    story.append(Paragraph(f"<b>Coverage Level:</b> {certificate.get_coverage_level_display()}", normal_style))
    story.append(Paragraph(f"<b>Insured Amount:</b> ${certificate.insured_amount:,.2f}", normal_style))
    story.append(Paragraph(f"<b>Premium Amount:</b> ${certificate.premium_amount:,.2f}", normal_style))
    story.append(Paragraph(f"<b>Total Premium:</b> ${certificate.total_premium:,.2f}", normal_style))
    
    # Insurance Period
    story.append(Paragraph("INSURANCE PERIOD", heading_style))
    story.append(Paragraph(f"<b>Start Date:</b> {certificate.start_date.strftime('%B %d, %Y')}", normal_style))
    story.append(Paragraph(f"<b>End Date:</b> {certificate.end_date.strftime('%B %d, %Y')}", normal_style))
    story.append(Paragraph(f"<b>Days Remaining:</b> {certificate.days_remaining} days", normal_style))
    
    # Financial Information
    story.append(Paragraph("FINANCIAL INFORMATION", heading_style))
    story.append(Paragraph(f"<b>Total Paid Amount:</b> ${certificate.total_paid_amount:,.2f}", normal_style))
    story.append(Paragraph(f"<b>Remaining Amount:</b> ${certificate.remaining_amount:,.2f}", normal_style))
    story.append(Paragraph(f"<b>Payment Status:</b> {certificate.get_payment_status_display()}", normal_style))
    
    # Agent Information
    if certificate.agent_name or certificate.agent_id:
        story.append(Paragraph("AGENT INFORMATION", heading_style))
        if certificate.agent_name:
            story.append(Paragraph(f"<b>Agent Name:</b> {certificate.agent_name}", normal_style))
        if certificate.agent_id:
            story.append(Paragraph(f"<b>Agent ID:</b> {certificate.agent_id}", normal_style))
    
    # Additional Information
    if certificate.description:
        story.append(Paragraph("DESCRIPTION", heading_style))
        story.append(Paragraph(certificate.description, normal_style))
    
    if certificate.terms_conditions:
        story.append(Paragraph("TERMS & CONDITIONS", heading_style))
        story.append(Paragraph(certificate.terms_conditions, normal_style))
    
    if certificate.special_conditions:
        story.append(Paragraph("SPECIAL CONDITIONS", heading_style))
        story.append(Paragraph(certificate.special_conditions, normal_style))
    
    # Certificate Status
    story.append(Paragraph("CERTIFICATE STATUS", heading_style))
    status_text = "ACTIVE" if certificate.is_active else "INACTIVE"
    status_color = colors.green if certificate.is_active else colors.red
    status_style = ParagraphStyle(
        'Status',
        parent=normal_style,
        textColor=status_color,
        fontSize=12,
        alignment=TA_CENTER
    )
    story.append(Paragraph(f"<b>Status: {status_text}</b>", status_style))
    
    # Footer
    story.append(Spacer(1, 30))
    footer_style = ParagraphStyle(
        'Footer',
        parent=normal_style,
        fontSize=8,
        alignment=TA_CENTER,
        textColor=colors.grey
    )
    story.append(Paragraph("This certificate is generated electronically by EverTrust Management System", footer_style))
    story.append(Paragraph(f"Generated on: {timezone.now().strftime('%B %d, %Y at %I:%M %p')}", footer_style))
    
    # Build PDF
    doc.build(story)
    
    # Get PDF content
    pdf_content = buffer.getvalue()
    buffer.close()
    
    return pdf_content

def fill_certificate_docx_template(certificate, template_path, output_path):
    doc = Document(template_path)
    # Replace placeholders in the DOCX template with certificate data
    mapping = {
        '{{policy_number}}': certificate.policy_number,
        '{{issue_date}}': certificate.issue_date.strftime('%d/%m/%Y') if certificate.issue_date else '',
        '{{client_name}}': certificate.client_name,
        '{{location}}': certificate.location,
        '{{activity}}': certificate.activity,
        '{{period_from}}': certificate.period_from.strftime('%d/%m/%Y') if certificate.period_from else '',
        '{{period_to}}': certificate.period_to.strftime('%d/%m/%Y') if certificate.period_to else '',
        '{{liability_per_person}}': f"{certificate.liability_per_person:,.2f} USD",
        '{{liability_per_occurrence}}': f"{certificate.liability_per_occurrence:,.2f} USD",
        '{{liability_aggregate}}': f"{certificate.liability_aggregate:,.2f} USD",
        '{{net_contribution}}': f"{certificate.net_contribution:,.2f}",
        '{{proportional_stamp}}': f"{certificate.proportional_stamp:,.2f}",
        '{{dimensional_stamp}}': f"{certificate.dimensional_stamp:,.2f}",
        '{{supervision_fees}}': f"{certificate.supervision_fees:,.2f}",
        '{{insurance_fees}}': f"{certificate.insurance_fees:,.2f}",
        '{{total_premium}}': f"{certificate.total_premium:,.2f}",
        '{{total_premium_words}}': certificate.total_premium_words,
        '{{coverage_level}}': certificate.coverage_level.title(),
        '{{certificate_title}}': certificate.certificate_title,
        '{{description_of_risk}}': "This policy indemnifies the Insured against legal liability for bodily injury to third parties arising from the Insured's business activities.",
        '{{client_bank_name}}': certificate.client_bank_name or '',
        '{{client_bank_account}}': certificate.client_bank_account or '',
    }
    for p in doc.paragraphs:
        for key, value in mapping.items():
            if key in p.text:
                p.text = p.text.replace(key, str(value))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in mapping.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(value))
    doc.save(output_path)

def docx_to_pdf(docx_path, pdf_path):
    """Convert DOCX file to PDF"""
    try:
        convert(docx_path, pdf_path)
        return True
    except Exception as e:
        print(f"Error converting DOCX to PDF: {e}")
        return False

def generate_certificate_from_template(certificate, template_path):
    """Generate certificate PDF from DOCX template"""
    try:
        # Prepare data for template
        data = {
            'POLICY_NUMBER': certificate.policy_number,
            'CLIENT_NAME': certificate.client_name,
            'CLIENT_EMAIL': certificate.client_email or 'N/A',
            'CLIENT_PHONE': certificate.client_phone,
            'CLIENT_ADDRESS': certificate.client_address,
            'CLIENT_ID_NUMBER': certificate.client_id_number,
            'CLIENT_DATE_OF_BIRTH': certificate.client_date_of_birth.strftime('%B %d, %Y'),
            'INSURANCE_TYPE': certificate.get_certificate_type_display(),
            'COVERAGE_LEVEL': certificate.get_coverage_level_display(),
            'INSURED_AMOUNT': f"${certificate.insured_amount:,.2f}",
            'PREMIUM_AMOUNT': f"${certificate.premium_amount:,.2f}",
            'DEDUCTIBLE_AMOUNT': f"${certificate.deductible_amount:,.2f}",

            'START_DATE': certificate.start_date.strftime('%B %d, %Y'),
            'END_DATE': certificate.end_date.strftime('%B %d, %Y'),
            'ISSUE_DATE': certificate.issue_date.strftime('%B %d, %Y'),
            'STATUS': certificate.get_status_display(),
            'PAYMENT_STATUS': certificate.get_payment_status_display(),
            'AGENT_NAME': certificate.agent_name or 'N/A',
            'AGENT_ID': certificate.agent_id or 'N/A',
            'DESCRIPTION': certificate.description or 'N/A',
            'TERMS_CONDITIONS': certificate.terms_conditions or 'N/A',
            'SPECIAL_CONDITIONS': certificate.special_conditions or 'N/A',
            'TOTAL_PAID': f"${certificate.total_paid_amount:,.2f}",
            'REMAINING_AMOUNT': f"${certificate.remaining_amount:,.2f}",
            'DAYS_REMAINING': f"{certificate.days_remaining} days" if certificate.days_remaining > 0 else "Expired",
            'GENERATED_DATE': timezone.now().strftime('%B %d, %Y at %I:%M %p'),
        }
        
        # Fill the template
        filled_doc = fill_certificate_docx_template(certificate, template_path, output_path)
        if not filled_doc:
            return None
        
        # Save filled document to temporary file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
            filled_doc.save(temp_docx.name)
            temp_docx_path = temp_docx.name
        
        # Convert to PDF
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
            temp_pdf_path = temp_pdf.name
        
        if docx_to_pdf(temp_docx_path, temp_pdf_path):
            # Read PDF content
            with open(temp_pdf_path, 'rb') as pdf_file:
                pdf_content = pdf_file.read()
            
            # Clean up temporary files
            os.unlink(temp_docx_path)
            os.unlink(temp_pdf_path)
            
            return pdf_content
        else:
            # Clean up temporary files
            if os.path.exists(temp_docx_path):
                os.unlink(temp_docx_path)
            return None
            
    except Exception as e:
        print(f"Error generating certificate from template: {e}")
        return None

def generate_certificate_html(certificate):
    """Generate HTML version of certificate"""
    context = {
        'certificate': certificate,
        'generated_date': timezone.now(),
    }
    
    return render_to_string('certificates/certificate_template.html', context)

def generate_simple_pdf(certificate):
    """Generate a simple PDF using ReportLab without WeasyPrint dependencies"""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from io import BytesIO
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        alignment=1,  # Center alignment
        textColor=colors.darkblue
    )
    story.append(Paragraph("INSURANCE CERTIFICATE", title_style))
    story.append(Spacer(1, 20))
    
    # Certificate Details
    data = [
        ['Policy Number:', certificate.policy_number],
        ['Certificate Type:', certificate.get_certificate_type_display()],
        ['Status:', certificate.get_status_display()],
        ['Client Name:', certificate.client_name],
        ['Client Email:', certificate.client_email],
        ['Client Phone:', certificate.client_phone],
        ['Insured Amount:', f"${certificate.insured_amount:,.2f}"],
        ['Premium Amount:', f"${certificate.premium_amount:,.2f}"],
        ['Start Date:', certificate.start_date.strftime('%B %d, %Y')],
        ['End Date:', certificate.end_date.strftime('%B %d, %Y')],
        ['Issue Date:', certificate.issue_date.strftime('%B %d, %Y')],
    ]
    
    # Create table
    table = Table(data, colWidths=[2*inch, 4*inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.lightblue),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 12),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 14),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    story.append(table)
    story.append(Spacer(1, 30))
    
    # Additional Information
    if certificate.description:
        story.append(Paragraph("Description:", styles['Heading2']))
        story.append(Paragraph(certificate.description, styles['Normal']))
        story.append(Spacer(1, 20))
    
    if certificate.terms_conditions:
        story.append(Paragraph("Terms & Conditions:", styles['Heading2']))
        story.append(Paragraph(certificate.terms_conditions, styles['Normal']))
        story.append(Spacer(1, 20))
    
    # Footer
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=10,
        alignment=1,  # Center alignment
        textColor=colors.grey
    )
    story.append(Spacer(1, 50))
    story.append(Paragraph("Generated by EverTrust Insurance System", footer_style))
    story.append(Paragraph(f"Certificate ID: {certificate.policy_number}", footer_style))
    
    # Build PDF
    doc.build(story)
    pdf_content = buffer.getvalue()
    buffer.close()
    
    return pdf_content

def generate_certificate_pdf_simple(certificate):
    """Generate a professional certificate PDF using only ReportLab"""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from io import BytesIO
    import os
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, 
                          rightMargin=2*cm, leftMargin=2*cm, 
                          topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    story = []
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        alignment=1,  # Center
        textColor=colors.darkblue,
        fontName='Helvetica-Bold'
    )
    
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Heading2'],
        fontSize=16,
        spaceAfter=20,
        alignment=1,  # Center
        textColor=colors.darkblue,
        fontName='Helvetica-Bold'
    )
    
    header_style = ParagraphStyle(
        'CustomHeader',
        parent=styles['Heading3'],
        fontSize=14,
        spaceAfter=10,
        textColor=colors.darkblue,
        fontName='Helvetica-Bold'
    )
    
    # Company Header
    story.append(Paragraph("EVERTRUST INSURANCE", title_style))
    story.append(Paragraph("Professional Insurance Certificate", subtitle_style))
    story.append(Spacer(1, 30))
    
    # Certificate Details
    story.append(Paragraph("Certificate Information", header_style))
    
    cert_data = [
        ['Policy Number:', certificate.policy_number],
        ['Certificate Type:', certificate.get_certificate_type_display()],
        ['Coverage Level:', certificate.get_coverage_level_display()],
        ['Status:', certificate.get_status_display()],
        ['Payment Status:', certificate.get_payment_status_display()],
    ]
    
    cert_table = Table(cert_data, colWidths=[2*inch, 4*inch])
    cert_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.lightblue),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    story.append(cert_table)
    story.append(Spacer(1, 20))
    
    # Client Information
    story.append(Paragraph("Client Information", header_style))
    
    client_data = [
        ['Client Name:', certificate.client_name],
        ['Client ID:', certificate.client_id_number],
        ['Date of Birth:', certificate.client_date_of_birth.strftime('%B %d, %Y')],
        ['Email:', certificate.client_email],
        ['Phone:', certificate.client_phone],
        ['Address:', certificate.client_address],
    ]
    
    if certificate.client_bank_name:
        client_data.append(['Bank Name:', certificate.client_bank_name])
    if certificate.client_bank_account:
        client_data.append(['Bank Account:', certificate.client_bank_account])
    
    client_table = Table(client_data, colWidths=[2*inch, 4*inch])
    client_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgreen),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('BACKGROUND', (0, 0), (-1, 0), colors.darkgreen),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    story.append(client_table)
    story.append(Spacer(1, 20))
    
    # Financial Information
    story.append(Paragraph("Financial Information", header_style))
    
    financial_data = [
        ['Insured Amount:', f"${certificate.insured_amount:,.2f}"],
        ['Premium Amount:', f"${certificate.premium_amount:,.2f}"],
        ['Total Premium:', f"${certificate.total_premium:,.2f}"],
        ['Net Contribution:', f"${certificate.net_contribution:,.2f}"],
        ['Proportional Stamp:', f"${certificate.proportional_stamp:,.2f}"],
        ['Dimensional Stamp:', f"${certificate.dimensional_stamp:,.2f}"],
        ['Supervision Fees:', f"${certificate.supervision_fees:,.2f}"],
        ['Insurance Fees:', f"${certificate.insurance_fees:,.2f}"],
    ]
    
    financial_table = Table(financial_data, colWidths=[2*inch, 4*inch])
    financial_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.lightyellow),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('BACKGROUND', (0, 0), (-1, 0), colors.orange),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    story.append(financial_table)
    story.append(Spacer(1, 20))
    
    # Coverage Period
    story.append(Paragraph("Coverage Period", header_style))
    
    period_data = [
        ['Start Date:', certificate.start_date.strftime('%B %d, %Y')],
        ['End Date:', certificate.end_date.strftime('%B %d, %Y')],
        ['Issue Date:', certificate.issue_date.strftime('%B %d, %Y')],
        ['Last Modified:', certificate.last_modified.strftime('%B %d, %Y')],
    ]
    
    period_table = Table(period_data, colWidths=[2*inch, 4*inch])
    period_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    story.append(period_table)
    story.append(Spacer(1, 20))
    
    # Additional Information
    if certificate.description:
        story.append(Paragraph("Description", header_style))
        story.append(Paragraph(certificate.description, styles['Normal']))
        story.append(Spacer(1, 15))
    
    if certificate.terms_conditions:
        story.append(Paragraph("Terms & Conditions", header_style))
        story.append(Paragraph(certificate.terms_conditions, styles['Normal']))
        story.append(Spacer(1, 15))
    
    if certificate.special_conditions:
        story.append(Paragraph("Special Conditions", header_style))
        story.append(Paragraph(certificate.special_conditions, styles['Normal']))
        story.append(Spacer(1, 15))
    
    # Footer
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=10,
        alignment=1,  # Center
        textColor=colors.grey
    )
    
    story.append(Spacer(1, 30))
    story.append(Paragraph("Generated by EverTrust Insurance System", footer_style))
    story.append(Paragraph(f"Certificate ID: {certificate.policy_number}", footer_style))
    story.append(Paragraph("For verification, contact: info@evertrust.uk", footer_style))
    
    # Build PDF
    doc.build(story)
    pdf_content = buffer.getvalue()
    buffer.close()
    
    return pdf_content 